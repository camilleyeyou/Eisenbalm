"""
build_editable_docx.py — Generate docs/client-editable/EISENBALM-EDITABLE-COPY.docx
from docs/client-editable/EISENBALM-EDITABLE-COPY.md using python-docx.

Run with:
    uv run --with python-docx python scripts/build_editable_docx.py

Re-runnable and idempotent — overwrites the .docx on each run.
Paths are resolved relative to the repository root (parent of this script's directory).
"""

# /// script
# dependencies = ["python-docx"]
# ///

import os
import sys

from docx import Document
from docx.shared import Pt


def repo_root() -> str:
    """Return the absolute path to the repository root (parent of scripts/)."""
    return os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def build_docx(md_path: str, docx_path: str) -> None:
    """Parse md_path and write a .docx to docx_path."""
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().splitlines()

    doc = Document()

    # Remove default empty paragraph Word adds
    for paragraph in doc.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)

    i = 0
    while i < len(lines):
        line = lines[i]

        # Heading levels
        if line.startswith("# ") and not line.startswith("## "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## ") and not line.startswith("### "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.strip() == "---":
            # Horizontal rule — add a short visual separator paragraph
            p = doc.add_paragraph()
            run = p.add_run("─" * 40)
            run.font.size = Pt(8)
            run.font.color.rgb = None  # let Word use default
        elif line.strip() == "":
            # Blank line — paragraph break (only add if previous wasn't already blank)
            doc.add_paragraph("")
        else:
            # Normal body paragraph
            doc.add_paragraph(line)

        i += 1

    os.makedirs(os.path.dirname(docx_path), exist_ok=True)
    doc.save(docx_path)
    print(f"Written: {docx_path}")


def main() -> None:
    root = repo_root()

    # Accept optional override paths as argv[1] and argv[2]
    if len(sys.argv) >= 3:
        md_path = sys.argv[1]
        docx_path = sys.argv[2]
    else:
        md_path = os.path.join(root, "docs", "client-editable", "EISENBALM-EDITABLE-COPY.md")
        docx_path = os.path.join(root, "docs", "client-editable", "EISENBALM-EDITABLE-COPY.docx")

    if not os.path.isfile(md_path):
        print(f"ERROR: Markdown master not found at {md_path}", file=sys.stderr)
        sys.exit(1)

    build_docx(md_path, docx_path)


if __name__ == "__main__":
    main()
